import os
import re
from PySide6.QtWidgets import QFileDialog, QMessageBox
from PySide6.QtCore import QThread, Signal, QTimer
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from config import MIN_MATCH_SCORE_FOR_REWRITE

# Import our internal modules
from utils.file_parser import FileParser
from api.groq_client import GroqAIClient
from db.db_manager import DBManager
from db.repository import UserRepository, ResumeRepository, JobHistoryRepository

class AIWorker(QThread):
    finished = Signal(dict, str) # result_dict, rewritten_text
    error = Signal(str)
    progress = Signal(int)

    def __init__(self, ai_client, resume_text, jd_text):
        super().__init__()
        self.ai_client = ai_client
        self.resume_text = resume_text
        self.jd_text = jd_text

    def run(self):
        try:
            self.progress.emit(20)
            # Call AI for Analysis
            result = self.ai_client.analyze_resume(self.resume_text, self.jd_text)
            self.progress.emit(50)
            
            if result:
                match_score = result.get('match_score', 0)

                # Only rewrite if match score >= MIN_MATCH_SCORE_FOR_REWRITE
                if match_score >= MIN_MATCH_SCORE_FOR_REWRITE:
                    rewritten_text = self.ai_client.rewrite_resume(
                        self.resume_text, self.jd_text, match_score=match_score
                    )
                    self.progress.emit(100)
                    self.finished.emit(result, rewritten_text or "")
                else:
                    # Score too low — send result but no rewritten text
                    self.progress.emit(100)
                    self.finished.emit(result, "")
            else:
                self.error.emit("The AI failed to process the request. Check your API key.")
        except Exception as e:
            self.error.emit(str(e))

class ChatWorker(QThread):
    finished = Signal(str)
    
    def __init__(self, ai_client, message_history):
        super().__init__()
        self.ai_client = ai_client
        self.message_history = message_history
        
    def run(self):
        response = self.ai_client.chat_with_assistant(self.message_history)
        self.finished.emit(response)

class AppController:
    def __init__(self, main_window, user_id):
        self.view = main_window
        self.user_id = user_id
        load_dotenv()
        
        self.db_manager = DBManager()
        
        self.user_repo = UserRepository(self.db_manager)
        self.resume_repo = ResumeRepository(self.db_manager)
        self.history_repo = JobHistoryRepository(self.db_manager)
        
        self.ai_client = GroqAIClient(api_key=os.getenv("GROQ_API_KEY"))
        
        self.parser = FileParser()
        
        # State Variables (Store data between different button clicks)
        self.current_resume_text = ""
        self.auto_version_name = ""
        self.rewritten_text = ""

        # Store vault data for click lookups
        self.vault_data = []

        # Skeleton Animation Timer
        self.skeleton_timer = QTimer()
        self.skeleton_timer.timeout.connect(self.animate_skeleton)
        self.skeleton_dots = 0

        # --- Connect UI Buttons to Logic ---
        self.view.upload_box.clicked.connect(self.handle_upload)
        self.view.btn_analyze.clicked.connect(self.handle_analyze)
        self.view.btn_download.clicked.connect(self.handle_download)

        # --- Connect Vault list and actions ---
        self.view.vault_list.itemClicked.connect(self.on_vault_item_clicked)
        self.view.btn_vault_delete.clicked.connect(self.on_vault_delete)
        self.view.btn_vault_export.clicked.connect(self.on_vault_export)

        # --- Connect Chat UI ---
        self.chat_history = []
        self.view.btn_chat_toggle.clicked.connect(self.toggle_chat)
        self.view.btn_chat_close.clicked.connect(self.toggle_chat)
        self.view.btn_chat_send.clicked.connect(self.handle_chat_send)
        self.view.chat_input.returnPressed.connect(self.handle_chat_send)

        self.load_history()
        self.load_vault()

    def animate_skeleton(self):
        self.skeleton_dots = (self.skeleton_dots + 1) % 4
        dots = "." * self.skeleton_dots
        self.view.skeleton_label.setText(f"Analyzing & Rewriting\nPlease wait{dots}")


    def handle_upload(self):
        """Handles the file selection and text extraction"""
        file_path, _ = QFileDialog.getOpenFileName(
            self.view, "Open Resume", "", "Documents (*.pdf *.docx *.txt)"
        )
        
        if file_path:
            text, error = self.parser.extract_text(file_path)
            if error:
                QMessageBox.critical(self.view, "Error", error)
            else: 
                self.current_resume_text = text
                # Update the button text to show the file was loaded
                self.view.upload_box.setText(f"{os.path.basename(file_path)} Uploaded")

    def handle_analyze(self):
        """The main logic: Analysis -> Progress Bar -> Rewrite -> Result"""
        jd_text = self.view.jd_input.toPlainText()
        
        if not self.current_resume_text or not jd_text:
            QMessageBox.warning(self.view, "Missing Data", "Please upload a resume and paste the Job Description!")
            return

        self.view.progress_bar.setVisible(True)
        self.view.progress_bar.setValue(0)
        self.view.btn_analyze.setEnabled(False) # Disable button while running
        self.view.btn_download.setVisible(False) # Hide download button from previous run
        self.view.keywords_container.setVisible(False) # Hide keywords from previous run
        
        self.view.results_area.setVisible(False)
        self.view.skeleton_label.setVisible(True)
        self.skeleton_timer.start(500)
        
        self.worker = AIWorker(self.ai_client, self.current_resume_text, jd_text)
        self.worker.progress.connect(self.view.progress_bar.setValue)
        self.worker.finished.connect(self.on_ai_finished)
        self.worker.error.connect(self.on_ai_error)
        self.worker.start()

    def on_ai_error(self, err_msg):
        self.skeleton_timer.stop()
        self.view.skeleton_label.setVisible(False)
        self.view.results_area.setVisible(True)
        QMessageBox.critical(self.view, "AI Error", err_msg)
        self.view.progress_bar.setVisible(False)
        self.view.btn_analyze.setEnabled(True)

    def on_ai_finished(self, result, rewritten_text):
        self.skeleton_timer.stop()
        self.view.skeleton_label.setVisible(False)
        self.view.results_area.setVisible(True)

        self.view.btn_analyze.setEnabled(True)
        self.view.progress_bar.setVisible(False)
        
        # Extract data from the result dictionary
        self.auto_version_name = result.get('suggested_version_name', 'Optimized_Resume')
        score = result.get('match_score', 0)

        # --- Missing Keywords Display ---
        missing_keywords = result.get('missing_keywords', [])
        if isinstance(missing_keywords, list) and missing_keywords:
            # Build colored tag chips
            tags_html = ""
            for kw in missing_keywords:
                tags_html += (
                    f'<span style="background-color: #45475a; color: #f38ba8; '
                    f'padding: 4px 10px; margin: 3px; border-radius: 12px; '
                    f'font-size: 12px; display: inline-block;">{kw}</span>  '
                )
            self.view.keywords_flow.setText(tags_html)
            self.view.keywords_container.setVisible(True)
        else:
            self.view.keywords_container.setVisible(False)

        # --- Suggestions ---
        suggestions = result.get('improvement_suggestions', [])
        if isinstance(suggestions, list):
            formatted_sugg = "\n".join([f"- {s}" for s in suggestions])
        elif isinstance(suggestions, str):
            formatted_sugg = suggestions
        else:
            formatted_sugg = "No suggestions available."

        # --- Check if score is below threshold ---
        if isinstance(score, (int, float)) and score < MIN_MATCH_SCORE_FOR_REWRITE:
            display_text = (
                f" Match Score: {score}%\n"
                f"_________________________________________________\n"
                f" MATCH TOO LOW, Resume Rewrite Blocked\n\n"
                f"Your resume does not match this Job Description enough (below 20%).\n"
                f"The AI cannot generate a meaningful optimized resume for this role.\n\n"
                f" WHAT YOU CAN DO:\n{formatted_sugg}\n\n"
                f"Consider gaining some of the missing skills or applying for roles\n"
                f"that better match your current experience."
            )
            self.view.results_area.setPlainText(display_text)
            self.view.btn_download.setVisible(False)
            self.rewritten_text = ""
        else:
            # Normal flow — score >= 20%
            display_text = (
                f" Match Score: {score}%\n"
                f" Suggested Version: {self.auto_version_name}\n"
                f"_________________________________________________\n"
                f" RECOMMENDATIONS:\n{formatted_sugg}"
            )
            self.view.results_area.setPlainText(display_text)
            
            self.rewritten_text = rewritten_text
            if self.rewritten_text:
                self.view.btn_download.setVisible(True)

        # Save to history regardless of score
        try:
            print(" Saving analysis to history...")
            self.history_repo.save_analysis(
                self.user_id,
                self.auto_version_name,
                self.view.jd_input.toPlainText(),
                result.get('match_score'),
                result.get('improvement_suggestions')
            )
            print("Saved Successfully!!!")
            self.load_history()
        except Exception as e :
            print(f"History Save FAILED: {e}")

    def handle_download(self):
        """Saves the result to MySQL and exports as a .docx file"""
        if not self.rewritten_text: 
            return
        
        try:
            self.resume_repo.save_resume(
                user_id=self.user_id, 
                version_name=self.auto_version_name, 
                content=self.rewritten_text
            )
            # Refresh the vault after saving
            self.load_vault()
        except Exception as e:
            print(f"Database Save Error: {e}")

        filename = f"{self.auto_version_name}.docx"
        path, _ = QFileDialog.getSaveFileName(self.view, "Save Resume", filename, "Word Files (*.docx)")
        
        if path:
            try:
                self._export_to_docx(self.rewritten_text, path)
                QMessageBox.information(self.view, "Success", f"Resume saved and exported as {filename}!")
            except Exception as e:
                QMessageBox.critical(self.view, "File Error", f"Could not save file: {e}")

    def load_history(self):
        print(f"debug: loading history for user_id:{self.user_id}")
        self.view.history_list.clear()

        try:
            history_data = self.history_repo.get_recent_history(self.user_id)
            print(f" DEBUG: raw data from DB: {history_data} ")
            print(f" DEBUG: number of entries found: {len(history_data) if history_data else 0} ")

            if not history_data:
                print(f" DEBUG: no data returned from db .")
                self.view.history_list.addItem('No recent works Found!')
                return

            for entry in history_data:
                if isinstance(entry, dict):
                    company = entry.get('company_name', 'Unknown Company')
                    score = entry.get('match_score', '0')
                else:
                    company = entry[0]
                    score = entry[1]

                display_text = f"{company} ({score}%)"

                self.view.history_list.addItem(display_text)
                print(f"DEBUG: added to sidebar :{display_text}")
        except Exception as e:
            print(f"DEBUG: Critical error in load_history : {e}")

    def load_vault(self):
        """Loads all saved optimized resumes into the Resume Vault tab"""
        self.view.vault_list.clear()
        self.vault_data = []

        try:
            resumes = self.resume_repo.get_all_resumes_by_user(self.user_id)

            if not resumes:
                self.view.vault_list.addItem("No saved resumes yet.")
                return

            for resume in resumes:
                if isinstance(resume, dict):
                    version = resume.get('version_name', 'Untitled')
                    created = resume.get('created_at', '')
                    self.vault_data.append(resume)
                    display = f" {version}  ({str(created)[:10]})"
                    self.view.vault_list.addItem(display)
        except Exception as e:
            print(f"DEBUG: Error loading vault: {e}")

    def on_vault_item_clicked(self, item):
        """When a resume in the vault is clicked, display its content in the results area"""
        row = self.view.vault_list.row(item)
        
        if row < len(self.vault_data):
            resume = self.vault_data[row]
            content = resume.get('content', 'No content available.')
            version = resume.get('version_name', 'Unknown')
            created = resume.get('created_at', '')

            display = (
                f" RESUME VAULT : {version}\n"
                f" Saved: {created}\n"
                f"_________________________________________\n\n"
                f"{content}"
            )
            self.view.results_area.setPlainText(display)

    def toggle_chat(self):
        """Delegate to the view's toggle (handles positioning), then set focus."""
        self.view.toggle_chat()
        if self.view.chat_widget.isVisible():
            self.view.chat_input.setFocus()
            
    def hide_chat(self):
        if self.view._chat_visible:
            self.view.toggle_chat()
        
    def render_chat_history(self, is_typing=False):
        html = ""
        for msg in self.chat_history:
            content = msg['content'].replace('\n', '<br>')
            if msg["role"] == "user":
                html += f"<p style='color:#a6e3a1;'><b>You:</b> {content}</p>"
            else:
                html += f"<p style='color:white;'><b>Alex:</b> {content}</p>"
                
        if is_typing:
            html += "<p style='color:#a6adc8;'><i>Alex is typing...</i></p>"
            
        self.view.chat_history_area.setHtml(html)
        scrollbar = self.view.chat_history_area.verticalScrollBar()
        scrollbar.setValue(scrollbar.maximum())
        
    def handle_chat_send(self):
        user_text = self.view.chat_input.text().strip()
        if not user_text:
            return
            
        self.view.chat_input.clear()
        self.chat_history.append({"role": "user", "content": user_text})
        self.render_chat_history(is_typing=True)
        
        self.view.chat_input.setEnabled(False)
        self.view.btn_chat_send.setEnabled(False)
        
        self.chat_worker = ChatWorker(self.ai_client, list(self.chat_history))
        self.chat_worker.finished.connect(self.on_chat_finished)
        self.chat_worker.start()
        
    def on_chat_finished(self, ai_response):
        self.chat_history.append({"role": "assistant", "content": ai_response})
        self.render_chat_history(is_typing=False)
        
        self.view.chat_input.setEnabled(True)
        self.view.btn_chat_send.setEnabled(True)
        self.view.chat_input.setFocus()

    def on_vault_delete(self):
        selected = self.view.vault_list.currentRow()
        if selected < 0 or selected >= len(self.vault_data):
            QMessageBox.warning(self.view, "Selection Error", "Please select a resume to delete.")
            return
            
        resume_id = self.vault_data[selected].get('resume_id')
        if resume_id:
            confirm = QMessageBox.question(self.view, "Confirm Delete", "Are you sure you want to delete this resume?", QMessageBox.Yes | QMessageBox.No)
            if confirm == QMessageBox.Yes:
                self.resume_repo.delete_resume(resume_id)
                self.load_vault()
                self.view.results_area.clear()

    def on_vault_export(self):
        selected = self.view.vault_list.currentRow()
        if selected < 0 or selected >= len(self.vault_data):
            QMessageBox.warning(self.view, "Selection Error", "Please select a resume to export.")
            return
            
        resume = self.vault_data[selected]
        content = resume.get('content', '')
        version = resume.get('version_name', 'Export')
        
        if not content:
            QMessageBox.warning(self.view, "Empty", "This resume is empty.")
            return
            
        filename = f"{version}.docx"
        path, _ = QFileDialog.getSaveFileName(self.view, "Export Resume", filename, "Word Files (*.docx)")
        
        if path:
            try:
                self._export_to_docx(content, path)
                QMessageBox.information(self.view, "Success", f"Resume exported as {filename}!")
            except Exception as e:
                QMessageBox.critical(self.view, "File Error", f"Could not export file: {e}")

    def _export_to_docx(self, content, path):
        """Helper method to format Markdown text into a professional Word document."""
        doc = Document()
        
        # Adjust margins for a standard resume look
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

        for line in content.split('\n'):
            line = line.strip()
            if not line:
                # Add a little spacing for empty lines instead of full paragraphs
                continue
                
            if line.startswith('# '):
                p = doc.add_heading(level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._add_formatted_runs(p, line[2:])
            elif line.startswith('## '):
                p = doc.add_heading(level=2)
                self._add_formatted_runs(p, line[3:])
            elif line.startswith('### '):
                p = doc.add_heading(level=3)
                self._add_formatted_runs(p, line[4:])
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(style='List Bullet')
                self._add_formatted_runs(p, line[2:])
            else:
                p = doc.add_paragraph()
                self._add_formatted_runs(p, line)
                
        doc.save(path)

    def _add_formatted_runs(self, paragraph, text):
        """Helper to parse **bold** text and add it to a python-docx paragraph."""
        parts = re.split(r'\*\*(.*?)\*\*', text)
        for i, part in enumerate(parts):
            if not part:
                continue
            run = paragraph.add_run(part)
            if i % 2 == 1:
                run.bold = True
